"""Local artifact export helpers for meeting delegate sessions."""

from __future__ import annotations

import importlib.util
import os
import re
import shutil
import subprocess
from pathlib import Path
from typing import Any

from .html_pdf_renderer import HTMLPDFRenderError, HTMLPDFRenderer


class ArtifactExportError(RuntimeError):
    """Raised when a local meeting artifact cannot be rendered."""


class MeetingArtifactExporter:
    def __init__(self) -> None:
        self._pandoc = self._resolve_binary(
            env_name="DELEGATE_PANDOC_PATH",
            command_name="pandoc",
            fallbacks=[
                r"C:\Users\jung\AppData\Local\Microsoft\WinGet\Packages\JohnMacFarlane.Pandoc_Microsoft.Winget.Source_8wekyb3d8bbwe\pandoc-3.9.0.1\pandoc.exe",
            ],
        )
        self._soffice = self._resolve_binary(
            env_name="DELEGATE_SOFFICE_PATH",
            command_name="soffice",
            fallbacks=[
                r"C:\Program Files\LibreOffice\program\soffice.exe",
                "/Applications/LibreOffice.app/Contents/MacOS/soffice",
            ],
        )
        self._timeout = float(os.getenv("DELEGATE_ARTIFACT_TIMEOUT_SECONDS", "90"))
        configured_reference = os.getenv("DELEGATE_REFERENCE_DOC_PATH", "").strip()
        default_reference = Path("doc/templates/meeting-summary-reference.docx")
        if configured_reference:
            self._reference_doc = Path(configured_reference)
        elif default_reference.exists():
            self._reference_doc = default_reference
        else:
            self._reference_doc = None
        self._html_renderer = HTMLPDFRenderer()

    @property
    def pdf_ready(self) -> bool:
        if self._selected_pdf_renderer() == "html":
            return self._html_renderer.pdf_ready
        return bool(self._pandoc and self._soffice)

    def readiness(self) -> dict[str, Any]:
        blocking_reasons: list[str] = []
        quality_notes: list[str] = []
        selected_renderer = self._selected_pdf_renderer()
        docx_pdf_ready = bool(self._pandoc and self._soffice)
        html_readiness = self._html_renderer.readiness()
        html_pdf_ready = bool(html_readiness.get("html_pdf_ready"))
        if selected_renderer == "docx":
            if not self._pandoc:
                blocking_reasons.append("pandoc is not installed for summary.docx export.")
            if not self._soffice:
                blocking_reasons.append("LibreOffice is not installed for summary.pdf export.")
        else:
            if not self._pandoc:
                quality_notes.append("pandoc is unavailable, so optional DOCX export cannot run.")
            if not self._soffice:
                quality_notes.append("LibreOffice is unavailable, so DOCX-first PDF export cannot run.")
        if selected_renderer == "html" and not html_pdf_ready:
            blocking_reasons.extend(str(item) for item in html_readiness.get("blocking_reasons") or [])
        docx_polish_available = importlib.util.find_spec("docx") is not None
        if not docx_polish_available:
            quality_notes.append(
                "python-docx is not installed, so the DOCX/PDF styling polish layer will be skipped."
            )
        reference_doc_ready = bool(self._reference_doc and self._reference_doc.exists())
        if self._reference_doc and not reference_doc_ready:
            quality_notes.append(
                f"Configured reference DOCX was not found: {self._reference_doc}"
            )
        return {
            "docx_ready": bool(self._pandoc),
            "pdf_ready": self.pdf_ready,
            "selected_pdf_renderer": selected_renderer,
            "docx_pdf_ready": docx_pdf_ready,
            "html_pdf_ready": html_pdf_ready,
            "html_pdf_browser_path": html_readiness.get("html_pdf_browser_path"),
            "renderer_profile_notes": "renderer_profile is an optional base hint; direct renderer_* surface controls are also supported.",
            "pandoc_path": self._pandoc,
            "soffice_path": self._soffice,
            "reference_doc_path": str(self._reference_doc.resolve()) if self._reference_doc else None,
            "reference_doc_ready": reference_doc_ready,
            "docx_polish_available": docx_polish_available,
            "blocking_reasons": blocking_reasons,
            "quality_notes": quality_notes,
        }

    def export_summary_bundle(
        self,
        summary_markdown_path: Path,
        *,
        renderer_profile: str = "default",
        briefing: dict[str, Any] | None = None,
        rendering_policy: dict[str, Any] | None = None,
        postprocess_requests: list[dict[str, Any]] | None = None,
    ) -> list[dict[str, str]]:
        if self._selected_pdf_renderer() == "html":
            if not summary_markdown_path.exists():
                return []
            if briefing is None:
                raise ArtifactExportError("HTML-first artifact export requires a briefing payload.")
            try:
                return self._html_renderer.render_summary_bundle(
                    summary_markdown_path,
                    briefing=dict(briefing),
                    rendering_policy=rendering_policy,
                    postprocess_requests=postprocess_requests,
                )
            except HTMLPDFRenderError as exc:
                raise ArtifactExportError(str(exc)) from exc

        exports: list[dict[str, str]] = []
        if not summary_markdown_path.exists():
            return exports
        if not self._pandoc:
            return exports

        docx_path = summary_markdown_path.with_suffix(".docx")
        self._run(
            self._pandoc_command(summary_markdown_path, docx_path),
            error_label="Pandoc summary export",
        )
        self._polish_summary_docx(
            docx_path,
            renderer_profile=renderer_profile,
            rendering_policy=rendering_policy,
            postprocess_requests=postprocess_requests,
        )
        if docx_path.exists():
            exports.append({"format": "docx", "path": str(docx_path)})

        if not self._soffice or not docx_path.exists():
            return exports

        out_dir = summary_markdown_path.parent
        self._run(
            [
                self._soffice,
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(out_dir),
                str(docx_path),
            ],
            error_label="LibreOffice PDF export",
        )

        generated_pdf = out_dir / f"{docx_path.stem}.pdf"
        target_pdf = summary_markdown_path.with_suffix(".pdf")
        if generated_pdf.exists():
            if generated_pdf != target_pdf:
                generated_pdf.replace(target_pdf)
            exports.append({"format": "pdf", "path": str(target_pdf)})
        return exports

    def _pandoc_command(self, summary_markdown_path: Path, docx_path: Path) -> list[str]:
        command = [
                self._pandoc,
                str(summary_markdown_path),
                "-o",
                str(docx_path),
            ]
        if self._reference_doc and self._reference_doc.exists():
            command.extend(["--reference-doc", str(self._reference_doc)])
        return command

    def _selected_pdf_renderer(self) -> str:
        raw = str(os.getenv("DELEGATE_MEETING_ARTIFACT_PDF_RENDERER", "html")).strip().lower()
        if raw == "html":
            return "html"
        return "docx"

    def _resolve_binary(self, *, env_name: str, command_name: str, fallbacks: list[str]) -> str | None:
        configured = os.getenv(env_name, "").strip()
        if configured and Path(configured).exists():
            return configured

        discovered = shutil.which(command_name)
        if discovered:
            return discovered

        for candidate in fallbacks:
            if Path(candidate).exists():
                return candidate
        return None

    def _run(self, args: list[str], *, error_label: str) -> None:
        try:
            completed = subprocess.run(
                args,
                capture_output=True,
                text=True,
                encoding="utf-8",
                errors="replace",
                timeout=self._timeout,
                check=False,
            )
        except OSError as exc:
            raise ArtifactExportError(f"{error_label} could not start: {exc}") from exc
        except subprocess.TimeoutExpired as exc:
            raise ArtifactExportError(f"{error_label} timed out after {self._timeout} seconds.") from exc

        if completed.returncode != 0:
            details = completed.stderr.strip() or completed.stdout.strip() or f"exit code {completed.returncode}"
            raise ArtifactExportError(f"{error_label} failed: {details}")

    def _polish_summary_docx(
        self,
        docx_path: Path,
        *,
        renderer_profile: str = "default",
        rendering_policy: dict[str, Any] | None = None,
        postprocess_requests: list[dict[str, Any]] | None = None,
    ) -> None:
        if not docx_path.exists():
            return
        try:
            from docx import Document
            from docx.enum.section import WD_SECTION_START
            from docx.enum.table import WD_TABLE_ALIGNMENT
            from docx.enum.table import WD_ALIGN_VERTICAL
            from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            from docx.shared import Inches, Pt, RGBColor
        except Exception:
            return

        document = Document(docx_path)
        profile = self._renderer_profile_settings(renderer_profile, rendering_policy=rendering_policy)
        design = self._renderer_design_settings(rendering_policy=rendering_policy, profile=profile)

        def set_rfonts(target: Any, *, ascii_name: str, east_asia_name: str | None = None) -> None:
            east_asia_name = east_asia_name or ascii_name
            target.font.name = ascii_name
            element = getattr(target, "_element", None)
            if element is None:
                element = getattr(target, "element", None)
            if element is None:
                return
            rpr = element.get_or_add_rPr()
            rfonts = rpr.rFonts
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rpr.append(rfonts)
            rfonts.set(qn("w:ascii"), ascii_name)
            rfonts.set(qn("w:hAnsi"), ascii_name)
            rfonts.set(qn("w:cs"), ascii_name)
            rfonts.set(qn("w:eastAsia"), east_asia_name)

        def set_run_font(run: Any, *, ascii_name: str, east_asia_name: str | None = None, size_pt: float | None = None, color_hex: str | None = None, bold: bool | None = None) -> None:
            east_asia_name = east_asia_name or ascii_name
            run.font.name = ascii_name
            if size_pt is not None:
                run.font.size = Pt(size_pt)
            if color_hex:
                run.font.color.rgb = RGBColor.from_string(color_hex)
            if bold is not None:
                run.font.bold = bold
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.rFonts
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rpr.append(rfonts)
            rfonts.set(qn("w:ascii"), ascii_name)
            rfonts.set(qn("w:hAnsi"), ascii_name)
            rfonts.set(qn("w:cs"), ascii_name)
            rfonts.set(qn("w:eastAsia"), east_asia_name)

        def set_run_shading(run: Any, fill_hex: str) -> None:
            fill = self._normalize_color_hex(fill_hex)
            if not fill:
                return
            rpr = run._element.get_or_add_rPr()
            shading = rpr.find(qn("w:shd"))
            if shading is None:
                shading = OxmlElement("w:shd")
                rpr.append(shading)
            shading.set(qn("w:val"), "clear")
            shading.set(qn("w:color"), "auto")
            shading.set(qn("w:fill"), fill)

        def set_paragraph_bottom_border(paragraph: Any, *, color_hex: str, size: int = 10, space: int = 6) -> None:
            p_pr = paragraph._element.get_or_add_pPr()
            borders = p_pr.find(qn("w:pBdr"))
            if borders is None:
                borders = OxmlElement("w:pBdr")
                p_pr.append(borders)
            bottom = borders.find(qn("w:bottom"))
            if bottom is None:
                bottom = OxmlElement("w:bottom")
                borders.append(bottom)
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), str(size))
            bottom.set(qn("w:space"), str(space))
            bottom.set(qn("w:color"), color_hex)

        def clear_paragraph_bottom_border(paragraph: Any) -> None:
            p_pr = paragraph._element.get_or_add_pPr()
            borders = p_pr.find(qn("w:pBdr"))
            if borders is not None:
                p_pr.remove(borders)

        def set_paragraph_shading(paragraph: Any, fill_hex: str) -> None:
            fill = self._normalize_color_hex(fill_hex)
            if not fill:
                return
            p_pr = paragraph._element.get_or_add_pPr()
            shading = p_pr.find(qn("w:shd"))
            if shading is None:
                shading = OxmlElement("w:shd")
                p_pr.append(shading)
            shading.set(qn("w:val"), "clear")
            shading.set(qn("w:color"), "auto")
            shading.set(qn("w:fill"), fill)

        def clear_paragraph_shading(paragraph: Any) -> None:
            p_pr = paragraph._element.get_or_add_pPr()
            shading = p_pr.find(qn("w:shd"))
            if shading is not None:
                p_pr.remove(shading)

        def set_cell_shading(cell: Any, fill_hex: str) -> None:
            tc_pr = cell._tc.get_or_add_tcPr()
            shading = tc_pr.find(qn("w:shd"))
            if shading is None:
                shading = OxmlElement("w:shd")
                tc_pr.append(shading)
            shading.set(qn("w:val"), "clear")
            shading.set(qn("w:color"), "auto")
            shading.set(qn("w:fill"), fill_hex)

        def set_cell_margins(cell: Any, *, top: int = 90, bottom: int = 90, left: int = 110, right: int = 110) -> None:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = tc_pr.find(qn("w:tcMar"))
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for side, value in {"top": top, "bottom": bottom, "left": left, "right": right}.items():
                node = tc_mar.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                    tc_mar.append(node)
                node.set(qn("w:w"), str(value))
                node.set(qn("w:type"), "dxa")

        def insert_table_before_paragraph(paragraph: Any, *, rows: int, cols: int) -> Any:
            table = document.add_table(rows=rows, cols=cols)
            paragraph._element.addprevious(table._element)
            return table

        def insert_table_after_paragraph(paragraph: Any, *, rows: int, cols: int) -> Any:
            table = document.add_table(rows=rows, cols=cols)
            paragraph._element.addnext(table._element)
            return table

        def insert_table_after_element(element: Any, *, rows: int, cols: int) -> Any:
            table = document.add_table(rows=rows, cols=cols)
            element.addnext(table._element)
            return table

        def remove_paragraph(paragraph: Any) -> None:
            parent = paragraph._element.getparent()
            if parent is not None:
                parent.remove(paragraph._element)

        def normalize_text(value: Any) -> str:
            return " ".join(str(value or "").strip().lower().split())

        def numeric_setting(value: Any, *, default: float, minimum: float, maximum: float) -> float:
            text = str(value or "").strip()
            if not text:
                return default
            try:
                parsed = float(text)
            except (TypeError, ValueError):
                return default
            if parsed < minimum:
                return minimum
            if parsed > maximum:
                return maximum
            return parsed

        def alignment_from_hint(value: Any) -> Any:
            hint = normalize_text(value)
            if hint in {"center", "centre", "centered", "center aligned", "middle"}:
                return WD_ALIGN_PARAGRAPH.CENTER
            if hint in {"right", "right aligned", "end"}:
                return WD_ALIGN_PARAGRAPH.RIGHT
            if hint in {"justify", "justified", "full width"}:
                return WD_ALIGN_PARAGRAPH.JUSTIFY
            return WD_ALIGN_PARAGRAPH.LEFT

        def append_postprocess_assets_to_cell(cell: Any, assets: list[dict[str, str]], *, image_width: float) -> None:
            for asset in assets:
                title = str(asset.get("title") or "").strip()
                caption = str(asset.get("caption") or "").strip()
                if title:
                    title_paragraph = cell.add_paragraph()
                    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    title_paragraph.paragraph_format.space_before = Pt(8)
                    title_paragraph.paragraph_format.space_after = Pt(4)
                    title_run = title_paragraph.add_run(title)
                    set_run_font(
                        title_run,
                        ascii_name=heading_font,
                        east_asia_name=heading_font,
                        size_pt=10.2,
                        color_hex=profile["heading3_color"],
                        bold=True,
                    )
                picture_paragraph = cell.add_paragraph()
                picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                picture_paragraph.paragraph_format.space_before = Pt(2)
                picture_paragraph.paragraph_format.space_after = Pt(4)
                try:
                    picture_run = picture_paragraph.add_run()
                    picture_run.add_picture(str(asset["path"]), width=Inches(image_width))
                except Exception:
                    remove_paragraph(picture_paragraph)
                    continue
                if caption:
                    caption_paragraph = cell.add_paragraph()
                    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption_paragraph.paragraph_format.space_before = Pt(0)
                    caption_paragraph.paragraph_format.space_after = Pt(6)
                    caption_run = caption_paragraph.add_run(caption)
                    set_run_font(
                        caption_run,
                        ascii_name=body_font,
                        east_asia_name=body_font,
                        size_pt=9.4,
                        color_hex=profile["muted_color"],
                        bold=False,
                    )

        def insert_postprocess_assets_before_paragraph(
            paragraph: Any,
            assets: list[dict[str, str]],
            *,
            fill_hex: str,
        ) -> None:
            if paragraph is None or not assets:
                return
            asset_table = insert_table_before_paragraph(paragraph, rows=1, cols=1)
            custom_table_elements.add(asset_table._element)
            asset_table.alignment = WD_TABLE_ALIGNMENT.LEFT
            asset_cell = asset_table.cell(0, 0)
            set_cell_margins(asset_cell, top=130, bottom=120, left=130, right=130)
            set_cell_shading(asset_cell, fill_hex)
            append_postprocess_assets_to_cell(asset_cell, assets, image_width=image_width)

        def insert_postprocess_assets_after_paragraph(
            paragraph: Any,
            assets: list[dict[str, str]],
            *,
            fill_hex: str,
        ) -> None:
            if paragraph is None or not assets:
                return
            asset_table = insert_table_after_paragraph(paragraph, rows=1, cols=1)
            custom_table_elements.add(asset_table._element)
            asset_table.alignment = WD_TABLE_ALIGNMENT.LEFT
            asset_cell = asset_table.cell(0, 0)
            set_cell_margins(asset_cell, top=130, bottom=120, left=130, right=130)
            set_cell_shading(asset_cell, fill_hex)
            append_postprocess_assets_to_cell(asset_cell, assets, image_width=image_width)

        def find_heading_paragraph(
            heading_text: str,
            *,
            style_name: str | None = "Heading 2",
        ) -> Any | None:
            target = normalize_text(heading_text)
            if not target:
                return None
            for paragraph in list(document.paragraphs):
                paragraph_style = paragraph.style.name if paragraph.style is not None else ""
                if style_name and paragraph_style != style_name:
                    continue
                paragraph_text = normalize_text(paragraph.text)
                if (
                    paragraph_text == target
                    or target in paragraph_text
                    or paragraph_text in target
                ):
                    return paragraph
            return None

        def find_next_heading_two_after(paragraph: Any) -> Any | None:
            if paragraph is None:
                return None
            paragraphs = list(document.paragraphs)
            target_element = getattr(paragraph, "_element", None)
            start_index = next(
                (
                    index
                    for index, candidate in enumerate(paragraphs)
                    if getattr(candidate, "_element", None) is target_element
                ),
                -1,
            )
            if start_index < 0:
                return None
            for candidate in paragraphs[start_index + 1:]:
                candidate_style = candidate.style.name if candidate.style is not None else ""
                if candidate_style == "Heading 2" and str(candidate.text or "").strip():
                    return candidate
            return None

        def find_block_tail_paragraph(paragraph: Any) -> Any | None:
            if paragraph is None:
                return None
            next_heading = find_next_heading_two_after(paragraph)
            paragraphs = list(document.paragraphs)
            target_element = getattr(paragraph, "_element", None)
            start_index = next(
                (
                    index
                    for index, candidate in enumerate(paragraphs)
                    if getattr(candidate, "_element", None) is target_element
                ),
                -1,
            )
            if start_index < 0:
                return None
            if next_heading is None:
                return paragraphs[-1] if paragraphs else paragraph
            next_element = getattr(next_heading, "_element", None)
            next_index = next(
                (
                    index
                    for index, candidate in enumerate(paragraphs)
                    if getattr(candidate, "_element", None) is next_element
                ),
                -1,
            )
            if next_index < 0:
                return paragraph
            if next_index <= start_index:
                return paragraph
            return paragraphs[next_index - 1]

        top_margin = numeric_setting(design.get("page_top_margin_inches"), default=0.72, minimum=0.3, maximum=2.0)
        bottom_margin = numeric_setting(design.get("page_bottom_margin_inches"), default=0.72, minimum=0.3, maximum=2.0)
        left_margin = numeric_setting(design.get("page_left_margin_inches"), default=0.82, minimum=0.4, maximum=2.25)
        right_margin = numeric_setting(design.get("page_right_margin_inches"), default=0.82, minimum=0.4, maximum=2.25)
        body_line_spacing = numeric_setting(design.get("body_line_spacing"), default=1.32, minimum=1.0, maximum=2.4)
        list_line_spacing = numeric_setting(design.get("list_line_spacing"), default=1.28, minimum=1.0, maximum=2.4)
        heading2_space_before = numeric_setting(design.get("heading2_space_before_pt"), default=16.0, minimum=0.0, maximum=48.0)
        heading2_space_after = numeric_setting(design.get("heading2_space_after_pt"), default=7.0, minimum=0.0, maximum=36.0)
        heading3_space_before = numeric_setting(design.get("heading3_space_before_pt"), default=11.0, minimum=0.0, maximum=36.0)
        heading3_space_after = numeric_setting(design.get("heading3_space_after_pt"), default=4.0, minimum=0.0, maximum=24.0)
        title_space_after = numeric_setting(design.get("title_space_after_pt"), default=20.0, minimum=0.0, maximum=48.0)
        title_divider_size = max(0, int(round(numeric_setting(design.get("title_divider_size"), default=0.0, minimum=0.0, maximum=24.0))))
        title_divider_space = max(0, int(round(numeric_setting(design.get("title_divider_space"), default=14.0, minimum=0.0, maximum=48.0))))
        custom_table_elements: set[Any] = set()

        for section in document.sections:
            section.start_type = WD_SECTION_START.CONTINUOUS
            section.top_margin = Inches(top_margin)
            section.bottom_margin = Inches(bottom_margin)
            section.left_margin = Inches(left_margin)
            section.right_margin = Inches(right_margin)

        title_font = str(profile.get("title_font") or "Noto Serif KR")
        heading_font = str(profile.get("heading_font") or "Noto Sans KR")
        body_font = str(profile.get("body_font") or "Noto Sans KR")

        for style_name, font_name, east_font, size_pt, color_hex, bold in [
            ("Heading 1", title_font, title_font, 22.0, profile["heading1_color"], True),
            ("Heading 2", heading_font, heading_font, 14.0, profile["heading2_color"], True),
            ("Heading 3", heading_font, heading_font, 12.3, profile["heading3_color"], True),
            ("Normal", body_font, body_font, 10.5, profile["body_color"], False),
            ("List Paragraph", body_font, body_font, 10.3, profile["body_color"], False),
            ("Subtitle", body_font, body_font, 9.6, profile["muted_color"], False),
            ("Title", title_font, title_font, 22.0, profile["heading1_color"], True),
            ("Table Grid", body_font, body_font, 10.0, profile["body_color"], False),
        ]:
            try:
                style = document.styles[style_name]
            except KeyError:
                continue
            set_rfonts(style, ascii_name=font_name, east_asia_name=east_font)
            style.font.size = Pt(size_pt)
            style.font.color.rgb = RGBColor.from_string(color_hex)
            style.font.bold = bold
            pf = style.paragraph_format
            if style_name == "Heading 1":
                pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
                pf.space_before = Pt(0)
                pf.space_after = Pt(title_space_after)
                pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                pf.line_spacing = 1.0
                pf.keep_with_next = True
            elif style_name == "Heading 2":
                pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
                pf.space_before = Pt(heading2_space_before)
                pf.space_after = Pt(heading2_space_after)
                pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                pf.line_spacing = 1.05
                pf.keep_with_next = True
            elif style_name == "Heading 3":
                pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
                pf.space_before = Pt(heading3_space_before)
                pf.space_after = Pt(heading3_space_after)
                pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                pf.line_spacing = 1.08
                pf.keep_with_next = True
            elif style_name in {"Normal", "List Paragraph", "Table Grid"}:
                pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                pf.line_spacing = body_line_spacing if style_name == "Normal" else list_line_spacing
                pf.space_after = Pt(4)

        title_paragraph: Any | None = document.paragraphs[0] if document.paragraphs else None
        if title_paragraph is not None:
            title = title_paragraph
            title_alignment = alignment_from_hint(design.get("cover_align"))
            kicker_text = str(design.get("cover_kicker") or "").strip()
            use_cover_panel = (
                bool(kicker_text)
                or bool(str(design.get("surface_tint") or "").strip())
                or bool(str(design.get("cover_fill") or "").strip())
            )
            if use_cover_panel and (title.text or "").strip():
                cover_table = insert_table_before_paragraph(title, rows=1, cols=1)
                custom_table_elements.add(cover_table._element)
                cover_table.alignment = (
                    WD_TABLE_ALIGNMENT.CENTER if title_alignment == WD_ALIGN_PARAGRAPH.CENTER else WD_TABLE_ALIGNMENT.LEFT
                )
                cover_cell = cover_table.cell(0, 0)
                set_cell_margins(
                    cover_cell,
                    top=180,
                    bottom=170,
                    left=180,
                    right=180,
                )
                set_cell_shading(cover_cell, str(design.get("cover_fill") or profile["table_label_fill"]))
                first_cover_paragraph = cover_cell.paragraphs[0]
                first_cover_paragraph.alignment = title_alignment
                first_cover_paragraph.paragraph_format.space_before = Pt(0)
                first_cover_paragraph.paragraph_format.space_after = Pt(0)
                if kicker_text:
                    kicker_run = first_cover_paragraph.add_run(f" {kicker_text} ")
                    set_run_font(
                        kicker_run,
                        ascii_name=heading_font,
                        east_asia_name=heading_font,
                        size_pt=8.3,
                        color_hex=str(design.get("kicker_text_color") or profile["body_color"]),
                        bold=True,
                    )
                    set_run_shading(kicker_run, str(design.get("kicker_fill") or profile["table_label_fill"]))
                    cover_title_paragraph = cover_cell.add_paragraph()
                else:
                    cover_title_paragraph = first_cover_paragraph
                cover_title_paragraph.alignment = title_alignment
                cover_title_paragraph.paragraph_format.space_before = Pt(10 if kicker_text else 0)
                cover_title_paragraph.paragraph_format.space_after = Pt(0)
                cover_title_paragraph.paragraph_format.keep_with_next = True
                title_run = cover_title_paragraph.add_run((title.text or "").strip())
                set_run_font(
                    title_run,
                    ascii_name=title_font,
                    east_asia_name=title_font,
                    size_pt=22.0,
                    color_hex=profile["heading1_color"],
                    bold=True,
                )
                remove_paragraph(title)
                title_paragraph = None
            else:
                if kicker_text:
                    kicker_paragraph = title.insert_paragraph_before()
                    kicker_paragraph.alignment = title_alignment
                    kicker_paragraph.paragraph_format.space_before = Pt(0)
                    kicker_paragraph.paragraph_format.space_after = Pt(9)
                    kicker_paragraph.paragraph_format.keep_with_next = True
                    kicker_run = kicker_paragraph.add_run(f" {kicker_text} ")
                    set_run_font(
                        kicker_run,
                        ascii_name=heading_font,
                        east_asia_name=heading_font,
                        size_pt=8.3,
                        color_hex=str(design.get("kicker_text_color") or profile["body_color"]),
                        bold=True,
                    )
                    set_run_shading(kicker_run, str(design.get("kicker_fill") or profile["table_label_fill"]))
                title.alignment = title_alignment
                title.paragraph_format.space_before = Pt(0)
                title.paragraph_format.space_after = Pt(title_space_after)
                title.paragraph_format.keep_with_next = True
                title_divider_color = self._normalize_color_hex(
                    design.get("title_divider_color")
                )
                if title_divider_color and title_divider_size > 0:
                    set_paragraph_bottom_border(
                        title,
                        color_hex=title_divider_color,
                        size=title_divider_size,
                        space=title_divider_space,
                    )
                else:
                    clear_paragraph_bottom_border(title)
                for run in title.runs:
                    set_run_font(
                        run,
                        ascii_name=title_font,
                        east_asia_name=title_font,
                        size_pt=22.0,
                        color_hex=profile["heading1_color"],
                        bold=True,
                    )

        policy = dict(rendering_policy or {})
        overview_heading = str(policy.get("overview_heading") or "회의 개요").strip() or "회의 개요"
        metadata_labels = tuple(
            f"{label}:"
            for label in (
                str(policy.get("overview_datetime_label") or "회의 일시").strip() or "회의 일시",
                str(policy.get("overview_author_label") or "작성 주체").strip() or "작성 주체",
                str(policy.get("overview_session_id_label") or "세션 ID").strip() or "세션 ID",
                str(policy.get("overview_participants_label") or "참석자").strip() or "참석자",
            )
        )
        metadata_label_names = {label.rstrip(":") for label in metadata_labels}
        raised_by_label = "제기자"
        speakers_label = "주요 화자"
        timestamps_label = "타임스탬프"
        section_emphasis_labels = (f"{raised_by_label}:", f"{speakers_label}:")
        section_timestamp_label = f"{timestamps_label}:"
        within_overview_block = False
        overview_heading_paragraph: Any | None = None
        overview_metadata_paragraphs: list[Any] = []
        overview_metadata_rows: list[tuple[str, str]] = []

        for paragraph in list(document.paragraphs):
            if title_paragraph is not None and paragraph == title_paragraph:
                continue
            text = (paragraph.text or "").strip()
            style_name = paragraph.style.name if paragraph.style is not None else ""
            if style_name == "Heading 2":
                within_overview_block = text == overview_heading
                if within_overview_block:
                    overview_heading_paragraph = paragraph
                if str(design.get("section_band_fill") or "").strip():
                    clear_paragraph_shading(paragraph)
                    clear_paragraph_bottom_border(paragraph)
                    set_paragraph_shading(paragraph, str(design.get("section_band_fill") or profile["table_label_fill"]))
                elif str(profile.get("section_border_color") or "").strip():
                    clear_paragraph_shading(paragraph)
                    clear_paragraph_bottom_border(paragraph)
                    set_paragraph_bottom_border(paragraph, color_hex=profile["section_border_color"], size=6, space=6)
                else:
                    clear_paragraph_shading(paragraph)
                    clear_paragraph_bottom_border(paragraph)
                for run in paragraph.runs:
                    set_run_font(
                        run,
                        ascii_name=heading_font,
                        east_asia_name=heading_font,
                        size_pt=14.0,
                        color_hex=profile["heading2_color"],
                        bold=True,
                    )
            elif style_name == "Heading 3":
                within_overview_block = False
                for run in paragraph.runs:
                    set_run_font(
                        run,
                        ascii_name=heading_font,
                        east_asia_name=heading_font,
                        size_pt=12.3,
                        color_hex=profile["heading3_color"],
                        bold=True,
                    )
            elif within_overview_block and any(text.startswith(label) for label in metadata_labels):
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(7)
                paragraph.paragraph_format.left_indent = Inches(0.02)
                if ":" in text:
                    raw_label, raw_value = text.split(":", 1)
                    label_text = raw_label.strip()
                    value_text = raw_value.strip()
                    if label_text and value_text:
                        overview_metadata_rows.append((label_text, value_text))
                        overview_metadata_paragraphs.append(paragraph)
                for run in paragraph.runs:
                    set_run_font(
                        run,
                        ascii_name=body_font,
                        east_asia_name=body_font,
                        size_pt=10.7,
                        color_hex=profile["body_color"],
                        bold=run.bold,
                    )
            elif text.startswith(section_emphasis_labels):
                within_overview_block = False
                paragraph.paragraph_format.left_indent = Inches(0.12)
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    set_run_font(
                        run,
                        ascii_name=heading_font,
                        east_asia_name=heading_font,
                        size_pt=9.8,
                        color_hex=profile["heading3_color"],
                        bold=True,
                    )
            elif text.startswith(section_timestamp_label):
                within_overview_block = False
                paragraph.paragraph_format.left_indent = Inches(0.12)
                paragraph.paragraph_format.space_after = Pt(6)
                for run in paragraph.runs:
                    set_run_font(
                        run,
                        ascii_name=body_font,
                        east_asia_name=body_font,
                        size_pt=9.6,
                        color_hex=profile["muted_color"],
                        bold=False,
                    )
            elif style_name == "Normal":
                if text:
                    within_overview_block = False
                for run in paragraph.runs:
                    set_run_font(
                        run,
                        ascii_name=body_font,
                        east_asia_name=body_font,
                        size_pt=10.5,
                        color_hex=profile["body_color"],
                        bold=run.bold,
                            )

        use_overview_panel = bool(str(design.get("overview_panel_fill") or "").strip())
        use_overview_table = (
            not use_overview_panel
            and (
                bool(str(design.get("overview_label_fill") or "").strip())
                or bool(str(design.get("overview_value_fill") or "").strip())
            )
        )
        if overview_heading_paragraph is not None and overview_metadata_rows and (use_overview_table or use_overview_panel):
            if use_overview_table:
                overview_table = insert_table_after_paragraph(
                    overview_heading_paragraph,
                    rows=len(overview_metadata_rows),
                    cols=2,
                )
                custom_table_elements.add(overview_table._element)
                overview_table.alignment = WD_TABLE_ALIGNMENT.LEFT
                for row_index, (label_text, value_text) in enumerate(overview_metadata_rows):
                    label_cell = overview_table.cell(row_index, 0)
                    value_cell = overview_table.cell(row_index, 1)
                    set_cell_margins(label_cell, top=120, bottom=120, left=120, right=120)
                    set_cell_margins(value_cell, top=120, bottom=120, left=120, right=120)
                    set_cell_shading(label_cell, str(design.get("overview_label_fill") or profile["table_label_fill"]))
                    set_cell_shading(value_cell, str(design.get("overview_value_fill") or "FFFFFF"))
                    label_paragraph = label_cell.paragraphs[0]
                    value_paragraph = value_cell.paragraphs[0]
                    label_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    value_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    label_run = label_paragraph.add_run(label_text)
                    value_run = value_paragraph.add_run(value_text)
                    set_run_font(
                        label_run,
                        ascii_name=heading_font,
                        east_asia_name=heading_font,
                        size_pt=9.7,
                        color_hex=profile["heading2_color"],
                        bold=True,
                    )
                    set_run_font(
                        value_run,
                        ascii_name=body_font,
                        east_asia_name=body_font,
                        size_pt=9.9,
                        color_hex=profile["body_color"],
                        bold=False,
                    )
            else:
                overview_table = insert_table_after_paragraph(
                    overview_heading_paragraph,
                    rows=1,
                    cols=1,
                )
                custom_table_elements.add(overview_table._element)
                overview_table.alignment = WD_TABLE_ALIGNMENT.LEFT
                overview_cell = overview_table.cell(0, 0)
                set_cell_margins(overview_cell, top=150, bottom=150, left=150, right=150)
                set_cell_shading(overview_cell, str(design.get("overview_panel_fill") or design.get("overview_value_fill") or "FFFFFF"))
                first_overview_paragraph = overview_cell.paragraphs[0]
                first_overview_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                first_overview_paragraph.paragraph_format.space_before = Pt(0)
                first_overview_paragraph.paragraph_format.space_after = Pt(6)
                for row_index, (label_text, value_text) in enumerate(overview_metadata_rows):
                    paragraph = first_overview_paragraph if row_index == 0 else overview_cell.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(5)
                    label_run = paragraph.add_run(f"{label_text}  ")
                    value_run = paragraph.add_run(value_text)
                    set_run_font(
                        label_run,
                        ascii_name=heading_font,
                        east_asia_name=heading_font,
                        size_pt=9.7,
                        color_hex=profile["heading2_color"],
                        bold=True,
                    )
                    set_run_font(
                        value_run,
                        ascii_name=body_font,
                        east_asia_name=body_font,
                        size_pt=9.9,
                        color_hex=profile["body_color"],
                        bold=False,
                    )
            for paragraph in overview_metadata_paragraphs:
                remove_paragraph(paragraph)

        postprocess_images = self._resolve_postprocess_image_assets(
            postprocess_requests,
            base_dir=docx_path.parent,
        )
        image_width = self._postprocess_image_width_inches(rendering_policy)
        block_asset_fill = str(design.get("section_panel_fill") or "FFFFFF")

        title_paragraph = next(
            (
                paragraph
                for paragraph in list(document.paragraphs)
                if (paragraph.style.name if paragraph.style is not None else "") == "Heading 1"
                and str(paragraph.text or "").strip()
            ),
            None,
        )

        top_level_heading_paragraphs = [
            paragraph
            for paragraph in list(document.paragraphs)
            if (paragraph.style.name if paragraph.style is not None else "") == "Heading 2"
            and str(paragraph.text or "").strip()
        ]

        def resolve_asset_relation(asset: dict[str, str]) -> str:
            raw_text = str(asset.get("placement_notes") or "").strip()
            normalized = normalize_text(raw_text)
            if any(token in normalized for token in ("appendix", "annex", "부록", "별첨", "끝", "마지막")):
                return "appendix"
            if any(token in normalized for token in ("title", "cover", "표지", "제목")):
                return "title"
            if any(token in normalized for token in ("before", "앞", "이전", "직전", "prelude")):
                return "before"
            if any(token in normalized for token in ("after", "뒤", "후", "다음", "following")):
                return "after"
            if any(token in normalized for token in ("inside", "inline", "within", "본문", "문서 안", "섹션 안", "안쪽")):
                return "inside"
            if str(asset.get("target_heading") or "").strip():
                return "inside"
            return "after"

        def resolve_asset_anchor(asset: dict[str, str]) -> tuple[str, Any | None]:
            target_heading = str(asset.get("target_heading") or "").strip()
            placement_notes = str(asset.get("placement_notes") or "").strip()
            if target_heading:
                heading_anchor = find_heading_paragraph(target_heading, style_name=None)
                if heading_anchor is not None:
                    return "heading", heading_anchor
            normalized_notes = normalize_text(placement_notes)
            if any(token in normalized_notes for token in ("title", "cover", "표지", "제목")):
                return "title", title_paragraph
            heading_anchor = find_heading_paragraph(placement_notes, style_name=None)
            if heading_anchor is not None:
                return "heading", heading_anchor
            for candidate in top_level_heading_paragraphs:
                normalized_heading = normalize_text(candidate.text)
                if normalized_heading and normalized_heading in normalized_notes:
                    return "heading", candidate
            return "", None

        title_before_assets: list[dict[str, str]] = []
        title_after_assets: list[dict[str, str]] = []
        before_section_assets: dict[str, list[dict[str, str]]] = {}
        inline_section_assets: dict[str, list[dict[str, str]]] = {}
        after_section_assets: dict[str, list[dict[str, str]]] = {}
        block_before_assets: dict[str, list[dict[str, str]]] = {}
        block_inside_assets: dict[str, list[dict[str, str]]] = {}
        block_after_assets: dict[str, list[dict[str, str]]] = {}
        appendix_assets: list[dict[str, str]] = []

        for asset in postprocess_images:
            relation = resolve_asset_relation(asset)
            anchor_kind, anchor = resolve_asset_anchor(asset)
            if relation == "appendix" or anchor is None:
                appendix_assets.append(asset)
                continue
            anchor_style = anchor.style.name if getattr(anchor, "style", None) is not None else ""
            if anchor_kind == "title" or anchor_style == "Heading 1":
                if relation == "before":
                    title_before_assets.append(asset)
                else:
                    title_after_assets.append(asset)
                continue
            if anchor_style == "Heading 3":
                heading_key = normalize_text(anchor.text)
                if relation == "before":
                    before_section_assets.setdefault(heading_key, []).append(asset)
                elif relation == "after":
                    after_section_assets.setdefault(heading_key, []).append(asset)
                else:
                    inline_section_assets.setdefault(heading_key, []).append(asset)
                continue
            block_key = normalize_text(anchor.text if anchor is not None else "")
            if not block_key:
                appendix_assets.append(asset)
                continue
            if relation == "before":
                block_before_assets.setdefault(block_key, []).append(asset)
            elif relation == "inside":
                block_inside_assets.setdefault(block_key, []).append(asset)
            else:
                block_after_assets.setdefault(block_key, []).append(asset)

        if title_paragraph is not None:
            if title_before_assets:
                insert_postprocess_assets_before_paragraph(
                    title_paragraph,
                    title_before_assets,
                    fill_hex=block_asset_fill,
                )
            if title_after_assets:
                insert_postprocess_assets_after_paragraph(
                    title_paragraph,
                    title_after_assets,
                    fill_hex=block_asset_fill,
                )
        else:
            appendix_assets.extend(title_before_assets)
            appendix_assets.extend(title_after_assets)

        for anchor_heading in top_level_heading_paragraphs:
            block_key = normalize_text(anchor_heading.text)
            if not block_key:
                continue
            before_assets = block_before_assets.pop(block_key, [])
            if before_assets:
                insert_postprocess_assets_before_paragraph(
                    anchor_heading,
                    before_assets,
                    fill_hex=block_asset_fill,
                )
            inside_assets = block_inside_assets.pop(block_key, [])
            if inside_assets:
                next_heading = find_next_heading_two_after(anchor_heading)
                if next_heading is not None:
                    insert_postprocess_assets_before_paragraph(
                        next_heading,
                        inside_assets,
                        fill_hex=block_asset_fill,
                    )
                else:
                    insert_postprocess_assets_after_paragraph(
                        find_block_tail_paragraph(anchor_heading),
                        inside_assets,
                        fill_hex=block_asset_fill,
                    )
            after_assets = block_after_assets.pop(block_key, [])
            if after_assets:
                next_heading = find_next_heading_two_after(anchor_heading)
                if next_heading is not None:
                    insert_postprocess_assets_before_paragraph(
                        next_heading,
                        after_assets,
                        fill_hex=block_asset_fill,
                    )
                else:
                    insert_postprocess_assets_after_paragraph(
                        find_block_tail_paragraph(anchor_heading),
                        after_assets,
                        fill_hex=block_asset_fill,
                    )

        for leftovers in (block_before_assets, block_inside_assets, block_after_assets):
            for assets in leftovers.values():
                appendix_assets.extend(assets)

        section_blocks: list[dict[str, Any]] = []
        current_block: dict[str, Any] | None = None
        for paragraph in list(document.paragraphs):
            style_name = paragraph.style.name if paragraph.style is not None else ""
            text = (paragraph.text or "").strip()
            if style_name == "Heading 3" and text:
                if current_block:
                    section_blocks.append(current_block)
                current_block = {"paragraphs": [paragraph], "heading_text": text}
                continue
            if current_block is None:
                continue
            if style_name in {"Heading 1", "Heading 2"} and text:
                section_blocks.append(current_block)
                current_block = None
                continue
            current_block["paragraphs"].append(paragraph)
        if current_block:
            section_blocks.append(current_block)

        use_section_panel = bool(str(design.get("section_panel_fill") or "").strip())
        use_section_accent = (not use_section_panel) and bool(str(design.get("section_accent_fill") or "").strip())
        if use_section_panel or use_section_accent:
            paragraph_elements = {paragraph._element for paragraph in document.paragraphs}
            for block in section_blocks:
                paragraphs = [
                    paragraph
                    for paragraph in list(block.get("paragraphs") or [])
                    if getattr(paragraph, "_element", None) in paragraph_elements
                ]
                if not paragraphs:
                    continue
                anchor = paragraphs[0]
                block_before_assets = before_section_assets.pop(normalize_text(block.get("heading_text") or ""), [])
                if block_before_assets:
                    insert_postprocess_assets_before_paragraph(
                        anchor,
                        block_before_assets,
                        fill_hex=block_asset_fill,
                    )
                if use_section_panel:
                    section_table = insert_table_before_paragraph(anchor, rows=1, cols=1)
                    custom_table_elements.add(section_table._element)
                    section_table.alignment = WD_TABLE_ALIGNMENT.LEFT
                    content_cell = section_table.cell(0, 0)
                    set_cell_margins(content_cell, top=150, bottom=140, left=150, right=150)
                    set_cell_shading(content_cell, str(design.get("section_panel_fill") or profile["table_label_fill"]))
                else:
                    section_table = insert_table_before_paragraph(anchor, rows=1, cols=2)
                    custom_table_elements.add(section_table._element)
                    section_table.alignment = WD_TABLE_ALIGNMENT.LEFT
                    try:
                        section_table.autofit = False
                    except Exception:
                        pass
                    accent_cell = section_table.cell(0, 0)
                    content_cell = section_table.cell(0, 1)
                    set_cell_margins(accent_cell, top=0, bottom=0, left=0, right=0)
                    set_cell_margins(content_cell, top=140, bottom=130, left=140, right=140)
                    set_cell_shading(accent_cell, str(design.get("section_accent_fill") or profile["heading2_color"]))
                    set_cell_shading(content_cell, "FFFFFF")
                    try:
                        section_table.columns[0].width = Inches(0.16)
                        section_table.columns[1].width = Inches(6.35)
                        accent_cell.width = Inches(0.16)
                        content_cell.width = Inches(6.35)
                    except Exception:
                        pass
                first_content_paragraph = content_cell.paragraphs[0]
                first_content_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                first_content_paragraph.paragraph_format.space_before = Pt(0)
                first_content_paragraph.paragraph_format.space_after = Pt(0)
                for idx, original in enumerate(paragraphs):
                    text = (original.text or "").strip()
                    if not text:
                        continue
                    style_name = original.style.name if original.style is not None else ""
                    target_paragraph = first_content_paragraph if idx == 0 else content_cell.add_paragraph()
                    target_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    target_paragraph.paragraph_format.space_before = Pt(0)
                    if style_name == "Heading 3":
                        target_paragraph.paragraph_format.space_after = Pt(6)
                        run = target_paragraph.add_run(text)
                        set_run_font(
                            run,
                            ascii_name=heading_font,
                            east_asia_name=heading_font,
                            size_pt=12.3,
                            color_hex=profile["heading3_color"],
                            bold=True,
                        )
                    elif text.startswith(section_emphasis_labels):
                        target_paragraph.paragraph_format.space_after = Pt(2)
                        run = target_paragraph.add_run(text)
                        set_run_font(
                            run,
                            ascii_name=heading_font,
                            east_asia_name=heading_font,
                            size_pt=9.8,
                            color_hex=profile["heading3_color"],
                            bold=True,
                        )
                    elif text.startswith(section_timestamp_label):
                        target_paragraph.paragraph_format.space_after = Pt(4)
                        run = target_paragraph.add_run(text)
                        set_run_font(
                            run,
                            ascii_name=body_font,
                            east_asia_name=body_font,
                            size_pt=9.6,
                            color_hex=profile["muted_color"],
                            bold=False,
                        )
                    else:
                        target_paragraph.paragraph_format.space_after = Pt(4)
                        run = target_paragraph.add_run(text)
                        set_run_font(
                            run,
                            ascii_name=body_font,
                            east_asia_name=body_font,
                            size_pt=10.3,
                            color_hex=profile["body_color"],
                            bold=False,
                        )
                block_assets = inline_section_assets.pop(normalize_text(block.get("heading_text") or ""), [])
                if block_assets:
                    append_postprocess_assets_to_cell(content_cell, block_assets, image_width=image_width)
                block_after_assets = after_section_assets.pop(normalize_text(block.get("heading_text") or ""), [])
                if block_after_assets:
                    append_postprocess_assets_to_cell(content_cell, block_after_assets, image_width=image_width)
                for paragraph in paragraphs:
                    remove_paragraph(paragraph)
        elif inline_section_assets or after_section_assets:
            for block in section_blocks:
                paragraphs = list(block.get("paragraphs") or [])
                if not paragraphs:
                    continue
                block_before_assets = before_section_assets.pop(normalize_text(block.get("heading_text") or ""), [])
                if block_before_assets:
                    insert_postprocess_assets_before_paragraph(
                        paragraphs[0],
                        block_before_assets,
                        fill_hex=block_asset_fill,
                    )
                block_assets = inline_section_assets.pop(normalize_text(block.get("heading_text") or ""), [])
                trailing_assets = after_section_assets.pop(normalize_text(block.get("heading_text") or ""), [])
                if block_assets:
                    asset_table = insert_table_after_paragraph(paragraphs[-1], rows=1, cols=1)
                    custom_table_elements.add(asset_table._element)
                    asset_table.alignment = WD_TABLE_ALIGNMENT.LEFT
                    asset_cell = asset_table.cell(0, 0)
                    set_cell_margins(asset_cell, top=130, bottom=120, left=130, right=130)
                    set_cell_shading(asset_cell, str(design.get("section_panel_fill") or "FFFFFF"))
                    append_postprocess_assets_to_cell(asset_cell, block_assets, image_width=image_width)
                    if trailing_assets:
                        trailing_table = insert_table_after_element(asset_table._element, rows=1, cols=1)
                        custom_table_elements.add(trailing_table._element)
                        trailing_table.alignment = WD_TABLE_ALIGNMENT.LEFT
                        trailing_cell = trailing_table.cell(0, 0)
                        set_cell_margins(trailing_cell, top=130, bottom=120, left=130, right=130)
                        set_cell_shading(trailing_cell, str(design.get("section_panel_fill") or "FFFFFF"))
                        append_postprocess_assets_to_cell(trailing_cell, trailing_assets, image_width=image_width)
                elif trailing_assets:
                    asset_table = insert_table_after_paragraph(paragraphs[-1], rows=1, cols=1)
                    custom_table_elements.add(asset_table._element)
                    asset_table.alignment = WD_TABLE_ALIGNMENT.LEFT
                    asset_cell = asset_table.cell(0, 0)
                    set_cell_margins(asset_cell, top=130, bottom=120, left=130, right=130)
                    set_cell_shading(asset_cell, str(design.get("section_panel_fill") or "FFFFFF"))
                    append_postprocess_assets_to_cell(asset_cell, trailing_assets, image_width=image_width)

        for leftovers in before_section_assets.values():
            appendix_assets.extend(list(leftovers))
        for leftovers in inline_section_assets.values():
            appendix_assets.extend(list(leftovers))
        for leftovers in after_section_assets.values():
            appendix_assets.extend(list(leftovers))

        if document.tables:
            first_table = document.tables[0]
            if first_table._element not in custom_table_elements:
                first_column_values = [(row.cells[0].text or "").strip() for row in first_table.rows[1:]]
                if first_column_values and metadata_label_names.issuperset(set(first_column_values)):
                    first_table._element.getparent().remove(first_table._element)

        for index, table in enumerate(document.tables):
            if table._element in custom_table_elements:
                continue
            table.style = "Table Grid"
            for row_index, row in enumerate(table.rows):
                for cell_index, cell in enumerate(row.cells):
                    set_cell_margins(cell)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    if row_index == 0:
                        set_cell_shading(cell, profile["table_header_fill"])
                    elif index == 0 and cell_index == 0:
                        set_cell_shading(cell, profile["table_label_fill"])
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.space_before = Pt(0)
                        paragraph.paragraph_format.space_after = Pt(0)
                        for run in paragraph.runs:
                            if row_index == 0:
                                set_run_font(run, ascii_name=body_font, east_asia_name=body_font, size_pt=10.0, color_hex="FFFFFF", bold=True)
                            elif cell_index == 0:
                                set_run_font(
                                    run,
                                    ascii_name=heading_font,
                                    east_asia_name=heading_font,
                                    size_pt=10.0,
                                    color_hex=profile["heading2_color"],
                                    bold=True,
                                )
                            else:
                                set_run_font(
                                    run,
                                    ascii_name=body_font,
                                    east_asia_name=body_font,
                                    size_pt=10.0,
                                    color_hex=profile["body_color"],
                                    bold=False,
                                )

        if appendix_assets:
            appendix_heading = str(
                dict(rendering_policy or {}).get("postprocess_requests_heading")
                or "후속 시각 자료"
            ).strip() or "후속 시각 자료"
            heading_paragraph = document.add_paragraph(style="Heading 2")
            heading_paragraph.paragraph_format.space_before = Pt(18)
            heading_paragraph.paragraph_format.space_after = Pt(7)
            heading_paragraph.paragraph_format.keep_with_next = True
            if str(design.get("section_band_fill") or "").strip():
                clear_paragraph_shading(heading_paragraph)
                clear_paragraph_bottom_border(heading_paragraph)
                set_paragraph_shading(
                    heading_paragraph,
                    str(design.get("section_band_fill") or profile["table_label_fill"]),
                )
            elif str(profile.get("section_border_color") or "").strip():
                clear_paragraph_shading(heading_paragraph)
                clear_paragraph_bottom_border(heading_paragraph)
                set_paragraph_bottom_border(
                    heading_paragraph,
                    color_hex=profile["section_border_color"],
                    size=6,
                    space=6,
                )
            else:
                clear_paragraph_shading(heading_paragraph)
                clear_paragraph_bottom_border(heading_paragraph)
            heading_run = heading_paragraph.add_run(appendix_heading)
            set_run_font(
                heading_run,
                ascii_name=heading_font,
                east_asia_name=heading_font,
                size_pt=14.0,
                color_hex=profile["heading2_color"],
                bold=True,
            )
            for asset in appendix_assets:
                title = str(asset.get("title") or "").strip()
                caption = str(asset.get("caption") or "").strip()
                if title:
                    title_paragraph = document.add_paragraph(style="Heading 3")
                    title_paragraph.paragraph_format.space_before = Pt(11)
                    title_paragraph.paragraph_format.space_after = Pt(4)
                    title_paragraph.paragraph_format.keep_with_next = True
                    title_run = title_paragraph.add_run(title)
                    set_run_font(
                        title_run,
                        ascii_name=heading_font,
                        east_asia_name=heading_font,
                        size_pt=12.3,
                        color_hex=profile["heading3_color"],
                        bold=True,
                    )
                try:
                    document.add_picture(str(asset["path"]), width=Inches(image_width))
                    picture_paragraph = document.paragraphs[-1]
                    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    picture_paragraph.paragraph_format.space_before = Pt(4)
                    picture_paragraph.paragraph_format.space_after = Pt(4)
                except Exception:
                    continue
                if caption:
                    caption_paragraph = document.add_paragraph()
                    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption_paragraph.paragraph_format.space_before = Pt(0)
                    caption_paragraph.paragraph_format.space_after = Pt(8)
                    caption_run = caption_paragraph.add_run(caption)
                    set_run_font(
                        caption_run,
                        ascii_name=body_font,
                        east_asia_name=body_font,
                        size_pt=9.6,
                        color_hex=profile["muted_color"],
                        bold=False,
                    )

        document.save(docx_path)

    def _resolve_postprocess_image_assets(
        self,
        postprocess_requests: list[dict[str, Any]] | None,
        *,
        base_dir: Path,
    ) -> list[dict[str, str]]:
        supported_suffixes = {".png", ".jpg", ".jpeg", ".webp", ".bmp"}
        assets: list[dict[str, str]] = []
        for item in list(postprocess_requests or []):
            if not isinstance(item, dict):
                continue
            raw_path = str(item.get("image_path") or "").strip()
            if not raw_path:
                continue
            candidate = Path(raw_path).expanduser()
            if not candidate.is_absolute():
                candidate = (base_dir / candidate).resolve()
            else:
                candidate = candidate.resolve()
            candidate = self._prefer_raw_visual_path(candidate)
            if not candidate.exists() or candidate.suffix.lower() not in supported_suffixes:
                continue
            assets.append(
                {
                    "path": str(candidate),
                    "title": str(item.get("title") or "").strip(),
                    "caption": str(item.get("caption") or item.get("instruction") or "").strip(),
                    "placement_notes": str(item.get("placement_notes") or ""),
                    "target_heading": str(item.get("target_heading") or ""),
                }
            )
        return assets

    def _prefer_raw_visual_path(self, path: Path) -> Path:
        match = re.match(r"^(?P<stem>.+)-card-(?P<index>\d+)(?P<suffix>\.[^.]+)$", path.name)
        if not match:
            return path
        raw_candidate = path.with_name(f"{match.group('stem')}{match.group('suffix')}")
        if raw_candidate.exists():
            return raw_candidate
        return path

    def _postprocess_image_width_inches(self, rendering_policy: dict[str, Any] | None) -> float:
        raw = dict(rendering_policy or {}).get("postprocess_image_width_inches")
        try:
            width = float(str(raw or "").strip())
        except (TypeError, ValueError):
            width = 5.9
        return min(max(width, 1.5), 7.2)

    def _renderer_profile_settings(
        self,
        renderer_profile: str,
        *,
        rendering_policy: dict[str, Any] | None = None,
    ) -> dict[str, str]:
        normalized = str(renderer_profile or "default").strip().lower()
        if normalized == "briefing":
            normalized = "default"
        profiles = {
            "default": {
                "heading1_color": "204F3D",
                "heading2_color": "244E3B",
                "heading3_color": "355E4B",
                "body_color": "1F2933",
                "muted_color": "6B7C74",
                "title_border_color": "C8D8D1",
                "section_border_color": "D9E7E1",
                "table_header_fill": "244E3B",
                "table_label_fill": "EEF4F1",
                "title_font": "Noto Serif KR",
                "heading_font": "Noto Sans KR",
                "body_font": "Noto Sans KR",
            },
            "formal": {
                "heading1_color": "27313A",
                "heading2_color": "30414F",
                "heading3_color": "425463",
                "body_color": "1F2933",
                "muted_color": "66727D",
                "title_border_color": "D2D8DE",
                "section_border_color": "E1E5E8",
                "table_header_fill": "30414F",
                "table_label_fill": "EEF1F4",
                "title_font": "Noto Serif KR",
                "heading_font": "Noto Sans KR",
                "body_font": "Noto Sans KR",
            },
            "report": {
                "heading1_color": "1E4E79",
                "heading2_color": "245D91",
                "heading3_color": "2F6FA8",
                "body_color": "1F2933",
                "muted_color": "5F6F82",
                "title_border_color": "C8D7E6",
                "section_border_color": "D7E3EE",
                "table_header_fill": "245D91",
                "table_label_fill": "EDF4FA",
                "title_font": "Noto Serif KR",
                "heading_font": "Noto Sans KR",
                "body_font": "Noto Sans KR",
            },
        }
        profile = dict(profiles.get(normalized) or profiles["default"])
        policy = dict(rendering_policy or {})
        primary = self._normalize_color_hex(policy.get("renderer_primary_color"))
        accent = self._normalize_color_hex(policy.get("renderer_accent_color"))
        neutral = self._normalize_color_hex(policy.get("renderer_neutral_color"))
        if primary:
            profile["heading1_color"] = primary
            profile["heading2_color"] = primary
            profile["title_border_color"] = self._mix_hex(primary, "FFFFFF", 0.76)
            profile["section_border_color"] = self._mix_hex(primary, "FFFFFF", 0.86)
            profile["table_header_fill"] = primary
            profile["table_label_fill"] = self._mix_hex(primary, "FFFFFF", 0.92)
        if accent:
            profile["heading3_color"] = accent
        if neutral:
            profile["body_color"] = neutral
            profile["muted_color"] = self._mix_hex(neutral, "FFFFFF", 0.24)
        title_font = str(policy.get("renderer_title_font") or "").strip()
        heading_font = str(policy.get("renderer_heading_font") or "").strip()
        body_font = str(policy.get("renderer_body_font") or "").strip()
        if title_font:
            profile["title_font"] = title_font
        if heading_font:
            profile["heading_font"] = heading_font
        if body_font:
            profile["body_font"] = body_font
        for policy_key, profile_key in (
            ("renderer_heading1_color", "heading1_color"),
            ("renderer_heading2_color", "heading2_color"),
            ("renderer_heading3_color", "heading3_color"),
            ("renderer_body_text_color", "body_color"),
            ("renderer_muted_text_color", "muted_color"),
            ("renderer_title_divider_color", "title_border_color"),
            ("renderer_section_border_color", "section_border_color"),
            ("renderer_table_header_fill_color", "table_header_fill"),
            ("renderer_table_label_fill_color", "table_label_fill"),
        ):
            override = self._normalize_color_hex(policy.get(policy_key))
            if override:
                profile[profile_key] = override
        return profile

    def _renderer_design_settings(
        self,
        *,
        rendering_policy: dict[str, Any] | None,
        profile: dict[str, str],
    ) -> dict[str, str]:
        policy = dict(rendering_policy or {})
        cover_align = str(policy.get("renderer_cover_align") or "").strip()
        primary = self._normalize_color_hex(policy.get("renderer_primary_color")) or str(profile.get("heading2_color") or "")
        surface_tint = self._normalize_color_hex(policy.get("renderer_surface_tint_color"))
        cover_fill = self._normalize_color_hex(policy.get("renderer_cover_fill_color"))
        if not cover_fill and surface_tint:
            cover_fill = surface_tint
        section_band_fill = self._normalize_color_hex(policy.get("renderer_section_band_fill_color"))
        section_panel_fill = self._normalize_color_hex(policy.get("renderer_section_panel_fill_color"))
        section_accent_fill = self._normalize_color_hex(policy.get("renderer_section_accent_fill_color"))
        overview_label_fill = self._normalize_color_hex(policy.get("renderer_overview_label_fill_color"))
        overview_value_fill = self._normalize_color_hex(policy.get("renderer_overview_value_fill_color"))
        overview_panel_fill = self._normalize_color_hex(policy.get("renderer_overview_panel_fill_color"))
        kicker_fill = self._normalize_color_hex(policy.get("renderer_kicker_fill_color"))
        if not kicker_fill:
            kicker_fill = primary or str(profile.get("table_header_fill") or "")
        kicker_text_color = self._normalize_color_hex(policy.get("renderer_kicker_text_color"))
        if not kicker_text_color:
            kicker_text_color = self._contrast_text_color(
                kicker_fill,
                dark_fallback=self._normalize_color_hex(profile.get("heading3_color")) or self._normalize_color_hex(profile.get("body_color")),
            )
        title_divider_color = self._normalize_color_hex(policy.get("renderer_title_divider_color"))
        return {
            "cover_align": cover_align,
            "surface_tint": surface_tint,
            "cover_kicker": str(policy.get("renderer_cover_kicker") or "").strip(),
            "kicker_fill": kicker_fill,
            "kicker_text_color": kicker_text_color,
            "cover_fill": cover_fill,
            "section_band_fill": section_band_fill,
            "section_panel_fill": section_panel_fill,
            "section_accent_fill": section_accent_fill,
            "overview_label_fill": overview_label_fill,
            "overview_value_fill": overview_value_fill,
            "overview_panel_fill": overview_panel_fill,
            "title_divider_color": title_divider_color,
            "page_top_margin_inches": str(policy.get("renderer_page_top_margin_inches") or "").strip(),
            "page_bottom_margin_inches": str(policy.get("renderer_page_bottom_margin_inches") or "").strip(),
            "page_left_margin_inches": str(policy.get("renderer_page_left_margin_inches") or "").strip(),
            "page_right_margin_inches": str(policy.get("renderer_page_right_margin_inches") or "").strip(),
            "body_line_spacing": str(policy.get("renderer_body_line_spacing") or "").strip(),
            "list_line_spacing": str(policy.get("renderer_list_line_spacing") or "").strip(),
            "heading2_space_before_pt": str(policy.get("renderer_heading2_space_before_pt") or "").strip(),
            "heading2_space_after_pt": str(policy.get("renderer_heading2_space_after_pt") or "").strip(),
            "heading3_space_before_pt": str(policy.get("renderer_heading3_space_before_pt") or "").strip(),
            "heading3_space_after_pt": str(policy.get("renderer_heading3_space_after_pt") or "").strip(),
            "title_space_after_pt": str(policy.get("renderer_title_space_after_pt") or "").strip(),
            "title_divider_size": str(policy.get("renderer_title_divider_size") or "").strip(),
            "title_divider_space": str(policy.get("renderer_title_divider_space") or "").strip(),
        }

    def _contrast_text_color(self, fill_hex: str, *, dark_fallback: str = "1F2933") -> str:
        fill = self._normalize_color_hex(fill_hex)
        if not fill:
            return dark_fallback
        red = int(fill[0:2], 16)
        green = int(fill[2:4], 16)
        blue = int(fill[4:6], 16)
        luminance = 0.299 * red + 0.587 * green + 0.114 * blue
        return dark_fallback if luminance >= 160 else "FFFFFF"

    def _normalize_color_hex(self, value: Any) -> str:
        text = str(value or "").strip().lstrip("#").upper()
        if len(text) == 3 and all(ch in "0123456789ABCDEF" for ch in text):
            text = "".join(ch * 2 for ch in text)
        if len(text) == 6 and all(ch in "0123456789ABCDEF" for ch in text):
            return text
        return ""

    def _mix_hex(self, source_hex: str, target_hex: str, ratio: float) -> str:
        source = self._normalize_color_hex(source_hex) or "000000"
        target = self._normalize_color_hex(target_hex) or "FFFFFF"
        ratio = min(max(float(ratio), 0.0), 1.0)
        mixed = []
        for offset in (0, 2, 4):
            src = int(source[offset:offset + 2], 16)
            dst = int(target[offset:offset + 2], 16)
            value = round(src * (1.0 - ratio) + dst * ratio)
            mixed.append(f"{value:02X}")
        return "".join(mixed)
